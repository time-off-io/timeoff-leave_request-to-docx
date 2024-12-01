import configparser
import os.path
import re
from base64 import b64encode
from datetime import datetime
from os import path
from pathlib import Path
from typing import Optional, Union

import requests
from docx import Document

leaves_to_show = None


class Execute:
    """
        Execute Paragraphs KeyWords Replace
        paragraph: docx paragraph
    """

    def __init__(self, paragraph):
        self.paragraph = paragraph

    def p_replace(self, x: int, key: str, value: str):
        """
        paragraph replace
        The reason why you do not replace the text in a paragraph directly is that it will cause the original format to
        change. Replacing the text in runs will not cause the original format to change
        :param x:       paragraph id
        :param key:     Keywords that need to be replaced
        :param value:   The replaced keywords
        :return:
        """
        # Gets the coordinate index values of all the characters in this paragraph [{run_index , char_index}]
        p_maps = [{"run": y, "char": z}
                  for y, run in enumerate(self.paragraph.runs)
                  for z, char in enumerate(list(run.text))]
        # Handle the number of times key occurs in this paragraph, and record the starting position in the list.
        # Here, while self.text.find(key) >= 0, the {"ab":"abc"} term will enter an endless loop
        # Takes a single paragraph as an independent body and gets an index list of key positions within the paragraph, or if the paragraph contains multiple keys, there are multiple index values
        k_idx = [s for s in range(len(self.paragraph.text)) if
                 self.paragraph.text.find(key, s, len(self.paragraph.text)) == s]
        for i, start_idx in enumerate(reversed(k_idx)):  # Reverse order iteration
            end_idx = start_idx + len(key)  # The end position of the keyword in this paragraph
            k_maps = p_maps[
                     start_idx:end_idx]  # Map Slice List A list of dictionaries for sections that contain keywords in a paragraph
            self.r_replace(k_maps, value)
            print(f"\t |Paragraph {x + 1: >3}, object {i + 1: >3} replaced successfully! | {key} ===> {value}")

    def r_replace(self, k_maps: list, value: str):
        """
        :param k_maps: The list of indexed dictionaries containing keywords，
                       e.g:[{"run":15, "char":3},{"run":15, "char":4},{"run":16, "char":0}]
        :param value:
        :return:
        Accept arguments, removing the characters in k_maps from back to front,
        leaving the first one to replace with value. Note: Must be removed in reverse order, otherwise
        the list length change will cause IndexError: string index out of range
        """
        for i, position in enumerate(reversed(k_maps), start=1):
            y, z = position["run"], position["char"]
            run: object = self.paragraph.runs[y]  # "k_maps" may contain multiple run ids, which need to be separated
            # Pit: Instead of the replace() method, str is converted to list after a single word to prevent run.text from making an error in some cases (e.g., a single run contains a duplicate word)
            thisrun = list(run.text)
            if i < len(k_maps):
                thisrun.pop(z)  # Deleting a corresponding word
            if i == len(
                    k_maps):  # The last iteration (first word), that is, the number of iterations is equal to the length of k_maps
                thisrun[z] = value  # Replace the word in the corresponding position with the new content
            run.text = ''.join(thisrun)  # Recover


class WordReplace:
    """
        file: Microsoft Office word file，only support .docx type file
    """

    def __init__(self, file):
        self.docx = Document(file)

    def body_content(self, replace_dict: dict):
        print("\t☺Processing keywords in the body...")
        for key, value in replace_dict.items():
            for x, paragraph in enumerate(self.docx.paragraphs):
                Execute(paragraph).p_replace(x, key, value)
        print("\t |Body keywords in the text are replaced!")

    def body_tables(self, replace_dict: dict):
        print("\t☺Processing keywords in the body'tables...")
        for key, value in replace_dict.items():
            for table in self.docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for x, paragraph in enumerate(cell.paragraphs):
                            Execute(paragraph).p_replace(x, key, value)
        print("\t |Body'tables keywords in the text are replaced!")

    def header_content(self, replace_dict: dict):
        print("\t☺Processing keywords in the header'body ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for x, paragraph in enumerate(section.header.paragraphs):
                    Execute(paragraph).p_replace(x, key, value)
        print("\t |Header'body keywords in the text are replaced!")

    def header_tables(self, replace_dict: dict):
        print("\t☺Processing keywords in the header'tables ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for x, paragraph in enumerate(cell.paragraphs):
                                Execute(paragraph).p_replace(x, key, value)
        print("\t |Header'tables keywords in the text are replaced!")

    def footer_content(self, replace_dict: dict):
        print("\t☺Processing keywords in the footer'body ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for x, paragraph in enumerate(section.footer.paragraphs):
                    Execute(paragraph).p_replace(x, key, value)
        print("\t |Footer'body keywords in the text are replaced!")

    def footer_tables(self, replace_dict: dict):
        print("\t☺Processing keywords in the footer'tables ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for x, paragraph in enumerate(cell.paragraphs):
                                Execute(paragraph).p_replace(x, key, value)
        print("\t |Footer'tables keywords in the text are replaced!")

    def save(self, filepath: str):
        """
        :param filepath: File saving path
        :return:
        """
        self.docx.save(filepath)

    @staticmethod
    def docx_list(dirPath):
        """
        :param dirPath:
        :return: List of docx files in the current directory
        """
        fileList = []
        for roots, dirs, files in os.walk(dirPath):
            for file in files:
                if file.endswith("docx") and file[0] != "~":  # Find the docx document and exclude temporary files
                    fileRoot = os.path.join(roots, file)
                    fileList.append(fileRoot)
        print("This directory finds a total of {0} related files!".format(len(fileList)))
        return fileList


def get_headers(config: dict) -> dict:
    """Returns the headers for the API calls"""
    username = config["API"]["username"]
    password = config["API"]["password"]
    token = b64encode(f"{username}:{password}".encode('utf-8')).decode("ascii")

    return {
        "Content-Type": "application/json",
        "Authorization": f"Basic {token}"
    }


def read_configuration() -> dict:
    """Reads the configuration from the config.ini file and validates it."""
    # Read file
    config = configparser.ConfigParser()
    ini_path = Path(path.dirname(__file__), "config.ini")
    config.read(ini_path, encoding="utf8")

    # Validate API values
    if config["API"]["api_base_url"] is None:
        print("ERROR: api_base does not exist in config.ini")
        exit(1)
    if config["API"]["username"] is None:
        print("ERROR: username does not exist in config.ini")
        exit(2)
    if config["API"]["password"] is None:
        print("ERROR: password does not exist in config.ini")
        exit(3)

    # Validate INPUT values
    global leaves_to_show
    # Count the past weeks to fetch the leaves from
    ask_for_latest_leaves_to_show = config["INPUT"]["ask_for_latest_leaves_to_show"]
    leaves_to_show = config["INPUT"]["default_latest_leaves_to_show"]
    if ask_for_latest_leaves_to_show != "0":
        leaves_to_show_set = input(f"Δώστε το πλήθος των αιτημάτων που θέλετε να εμφανιστούν (πx. 10)? "
                                   f"default: {leaves_to_show} > ")
        leaves_to_show = leaves_to_show_set if leaves_to_show_set != "" else leaves_to_show
    if config["INPUT"]["leave_status"] is None:
        print("ERROR: leave_status does not exist in config.ini")
        exit(7)

    # Validate OUTPUT values
    template_dir = config["OUTPUT"]["template_dir"]
    if template_dir is None or not os.path.isdir(template_dir):
        print("ERROR: template_dir does not exist in the config.ini, or path is not a valid directory")
        exit(4)
    output_dir = config["OUTPUT"]["output_dir"]
    if output_dir is None or not os.path.isdir(output_dir):
        print("ERROR: output_dir does not exist in the config.ini, or path is not a valid directory")
        exit(5)
    if config["OUTPUT"]["date_format"] is None:
        print("ERROR: date_format does not exist in the config.ini")
        exit(6)
    if config["OUTPUT"]["filename_pattern"] is None:
        print("ERROR: filename_pattern does not exist in the config.ini")
        exit(8)

    return config


def make_request(config: dict, url_path: str) -> Union[list, dict]:
    """Makes a request to the API and returns the response as a list."""
    # Make the request
    url = f"{config['API']['api_base_url']}{url_path}"
    res = requests.get(url, headers=get_headers(config))

    # Check if the request was successful
    if res.status_code != 200:
        print(f"ERROR: url: {url}, response status: {res.status_code}, body: {res.text}")
        exit(50)

    # Convert the response to json
    return res.json()


def fetch_leaves_types(config: dict) -> list:
    """Fetches the leaves types from the API and returns them as a list."""
    return make_request(config, "/mod-leaves/leave-type/")


def fetch_employees(config: dict) -> list:
    """Fetches the employees from the API and returns them as a list."""
    return make_request(config, "/mod-personnel/employee/")


def fetch_leaves(config: dict, employees: list, leaves_types: list) -> list:
    """Fetches the leaves from the API and returns them as a list."""
    # Make the request
    leaves = make_request(config, "/mod-leaves/leave/")

    # Filter: keep the latest config["INPUT"]["latest_weeks_to_show"] leaves
    if len(leaves) > int(leaves_to_show):
        leaves = [leave for leave in leaves if leave["status"] == config["INPUT"]["leave_status"]]
        leaves = leaves[:int(leaves_to_show)]

    # Normalize the leaves
    for leave in leaves:
        leave["employee"] = [x for x in employees if x["id"] == leave["employeeId"]][0]
        leave["leaveType"] = [x for x in leaves_types if x["id"] == leave["leaveTypeId"]][0]
        leave["requestDate"] = datetime.strptime(leave["requestDate"], "%d.%m.%Y")
        leave["protocolNo"] = leave["authLevel1_reason"] if "authLevel1_reason" in leave else "-"
        leave["startDate"] = datetime.strptime(leave["startDate"], "%d.%m.%Y")
        leave["endDate"] = datetime.strptime(leave["endDate"], "%d.%m.%Y")

    # Sort the leaves by startDate
    leaves.sort(key=lambda x: x["startDate"], reverse=True)

    return leaves


def str_upto(value: str, up_to: int, fill_with: str = " ") -> str:
    """Returns a string with the length of up_to, filled with the value of fill_with variable."""
    if len(value) == up_to:
        return value

    if len(value) > up_to:
        return value[:up_to]

    while len(value) < up_to:
        value = f"{value}{fill_with}"
    return value


def build_table(leaves: list) -> None:
    """Prints a table with the leaves."""
    index = 0
    leave_type_title = str_upto('Τύπος Άδειας', 40)
    leave_type_fill = str_upto('', 40, fill_with='-')

    # Header
    print("Αιτήματα άδειας")
    print(str_upto("", 136, "-"))
    print(
        f"A/A  | Αριθμ. Πρωτοκόλλου | {leave_type_title} | Ημ/νία Έναρξης | Ημ/νία Λήξης | Κατάσταση | Επώνυμο/Όνομα")
    print(
        f"-----|--------------------|-{leave_type_fill}-|----------------|--------------|-----------|---------------")

    # Body
    for leave in leaves:
        index += 1
        print(f" {str(index) + ' ' if index < 10 else str(index)}  |"
              f" {str_upto(leave['protocolNo'].strip(), 18)} |"
              f" {str_upto(leave['leaveType']['title'], 40)} |"
              f" {leave['startDate'].strftime('%d.%m.%Y')}     |"
              f" {leave['endDate'].strftime('%d.%m.%Y')}   |"
              f" {str_upto(leave['status'], 9)} |"
              f" {leave['employee']['firstName']} {leave['employee']['lastName']}")

    print(str_upto("", 137, "-"))


def variables_to_replace(leave: dict, leave_got: dict, department_got: dict, config: dict) -> dict:
    """Returns a dictionary with the variables to replace in the template file."""
    return {
        "${TODAY}": datetime.now().strftime(config["OUTPUT"]["date_format"]),
        "${FIRSTNAME}": leave["employee"]["firstName"],
        "${LASTNAME}": leave["employee"]["lastName"],
        "${DEPARTMENT}": department_got["title"],
        "${LEAVE_TYPE}": leave["leaveType"]["title"],
        "${REQUEST_DATE}": leave["requestDate"].strftime(config["OUTPUT"]["date_format"]),
        "${START_DATE}": leave["startDate"].strftime(config["OUTPUT"]["date_format"]),
        "${END_DATE}": leave["endDate"].strftime(config["OUTPUT"]["date_format"]),
        "${DAYS_COUNT}": str(len(leave_got["leaveDays"])),
        "${REASON}": leave["remark"],
        "${REASON_AUTH1}": leave["authLevel1_reason"] if "authLevel1_reason" in leave else "",
        "${REASON_AUTH2}": leave["authLevel2_reason"] if "authLevel2_reason" in leave else "",
    }


def export_file_path(variables: dict, config: dict) -> Optional[str]:
    """Builds the filename based on the filename_pattern setting and returns the output file path for the document."""
    # The output file path pattern
    output_filename = config["OUTPUT"]["filename_pattern"]

    # Replace the placeholders with the variables values
    for key in variables:
        output_filename = output_filename.replace(key, variables[key])

    # Remove invalid characters from the output file path
    output_filename = re.sub(r'[^\w\s-]', '', output_filename.lower())

    # Validate the generated filename
    if len(output_filename) > 255:
        print(f"ERROR: Το αρχείο δημιουργήθηκε αλλά το όνομα του αρχείου docx είναι πολύ μεγάλο. "
              f"Αλλάξτε τη σχετική ρύθμιση του config.ini αρχείου!\n"
              f"Το όνομα του αρχείου είναι το: {output_filename}")
        return None
    if len(output_filename) == 0:
        print("ERROR: Το αρχείο δημιουργήθηκε αλλά το όνομα του αρχείου docx είναι κενό."
              f"Αλλάξτε την σχετική ρύθμιση του config.ini αρχείου!")
        return None

    return str(Path(config["OUTPUT"]["output_dir"], f"{output_filename}.docx"))


def export_document(config: dict, leave: dict) -> Optional[str]:
    """Exports the leave to a document"""
    # Get the valid template file path
    filename = leave['leaveType']["remark"]
    if filename == "":
        print(f"ERROR: Δεν έχετε ορίσει το όνομα του αρχείου για τον τύπο άδειας '{leave['leaveType']['title']}'. "
              f"Θα πρέπει να συνδεθείτε στο app.timeoff.gr, να πάτε στο μενού 'Άδειες > Τύποι αδειών' και, με "
              f"επεξεργασία του σχετικού τύπου άδειας, να ορίσετε το όνομα του template αρχείου στο πεδίο 'Περιγραφή'.")
        return None
    employment_status = leave["employee"]["vatNo"]
    template_file_path = str(Path(config["OUTPUT"]["template_dir"] + employment_status, filename))
    if not os.path.isfile(template_file_path):
        print(f"ERROR: Template file {template_file_path} does not exist!")
        return None

    # Get the count of leave days for this request and the department
    leave_got = make_request(config, f"/mod-leaves/leave/{leave['id']}")
    department_got = make_request(config, f"/mod-personnel/department/{leave['employee']['departmentId']}")

    # Prepare the variables to be used to replace the placeholders
    variables = variables_to_replace(leave, leave_got, department_got, config)

    # Replace the placeholders with the variables values
    word_replace = WordReplace(template_file_path)
    word_replace.header_content(variables)
    word_replace.header_tables(variables)
    word_replace.body_content(variables)
    word_replace.body_tables(variables)
    word_replace.footer_content(variables)
    word_replace.footer_tables(variables)

    # Get the output file path
    output_file = export_file_path(variables, config)

    # Save the document
    word_replace.docx.save(output_file)

    return output_file


def run() -> None:
    """Runs the script"""
    # Read configuration
    config = read_configuration()

    # Fetch data from the API
    employees = fetch_employees(config)
    leave_types = fetch_leaves_types(config)
    leaves = fetch_leaves(config, employees, leave_types)

    # Display leaves on a table
    build_table(leaves)

    # Request for the leave to be selected
    while True:
        leave_no = input("Γράψτε το Α/Α της αίτησης που θέλετε να εξάγετε? Για έξοδο γράψτε μηδέν (0) > ")
        if leave_no == "0":
            exit(0)

        # Validate the input
        try:
            leave_no = int(leave_no)
            if len(leaves) < leave_no or 0 > leave_no:
                print("ERROR: Ο αριθμός που δώσατε δεν υπάρχει στον πίνακα!")
                continue
        except ValueError:
            print("ERROR: Λάθος τιμή! Η τιμή που δώσατε δεν είναι αριθμός! ")
            continue

        # Export the leave
        filepath = export_document(config, leaves[leave_no - 1])
        if filepath is not None:
            print(f"Η δημιουργία του αρχείου ολοκληρώθηκε. Το αρχείο βρίσκεται στην παρακάτω διαδρομή:\n{filepath}\n")


if __name__ == "__main__":
    run()
